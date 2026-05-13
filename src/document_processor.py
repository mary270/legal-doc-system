"""
Document processor: handles messy legal PDFs and text files via OCR and text extraction.
Produces clean text + structured fields ready for downstream retrieval and drafting.
Uses Groq (LLaMA 3) for structured field extraction.
"""
from __future__ import annotations

import json
import os
import re
import uuid
from pathlib import Path
from typing import Any

from groq import Groq

# Optional heavy deps — gracefully degrade when not installed
try:
    import pdfplumber
    _PDF_OK = True
except ImportError:
    _PDF_OK = False

try:
    from PIL import Image
    import pytesseract
    _OCR_OK = True
except ImportError:
    _OCR_OK = False


_client = Groq(api_key=os.environ.get("GROQ_API_KEY", ""))
_MODEL = "llama-3.3-70b-versatile"

_STRUCT_PROMPT = """You are a legal document analyst.

Given the raw extracted text from a legal document, extract structured fields.
Return ONLY valid JSON with these fields (use null for missing fields):

{
  "document_type": "contract | deed | notice | memo | affidavit | court_filing | other",
  "parties": ["list of party names"],
  "dates": {
    "document_date": "YYYY-MM-DD or null",
    "effective_date": "YYYY-MM-DD or null",
    "other_dates": ["label: date string pairs"]
  },
  "property_or_subject": "brief description of the subject matter",
  "key_obligations": ["list of main obligations or terms"],
  "jurisdiction": "jurisdiction or null",
  "case_or_reference_number": "reference number or null",
  "summary": "1-2 sentence plain-English summary of what this document is about"
}

Raw document text:
"""


def _extract_text_pdf(file_path: str) -> str:
    """Try pdfplumber first; fall back to page-by-page OCR if text is sparse."""
    if not _PDF_OK:
        raise RuntimeError("pdfplumber not installed — cannot process PDFs")

    text_parts: list[str] = []
    ocr_pages: list[int] = []

    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if len(page_text.strip()) < 50:
                ocr_pages.append(i)
            else:
                text_parts.append(f"[Page {i + 1}]\n{page_text}")

        if ocr_pages and _OCR_OK:
            for i in ocr_pages:
                page = pdf.pages[i]
                img = page.to_image(resolution=200).original
                ocr_text = pytesseract.image_to_string(img)
                text_parts.insert(i, f"[Page {i + 1} — OCR]\n{ocr_text}")
        elif ocr_pages:
            for i in ocr_pages:
                text_parts.insert(i, f"[Page {i + 1} — OCR unavailable, text sparse]")

    return "\n\n".join(text_parts)


def _extract_text_image(file_path: str) -> str:
    if not _OCR_OK:
        raise RuntimeError("pytesseract/Pillow not installed — cannot OCR images")
    img = Image.open(file_path)
    return pytesseract.image_to_string(img)


def _extract_text_plain(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


def _extract_text_docx(file_path: str) -> str:
    """Extract text from .docx Word documents."""
    try:
        import docx
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab table content
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except ImportError:
        raise RuntimeError("python-docx not installed — run: pip install python-docx")


def _extract_text_rtf(file_path: str) -> str:
    """Extract text from RTF files by stripping RTF control words."""
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    # Strip RTF control sequences
    text = re.sub(r"\\\w+", " ", raw)
    text = re.sub(r"[{}]", "", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _extract_text_csv(file_path: str) -> str:
    """Convert CSV to readable text representation."""
    import csv
    rows = []
    with open(file_path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for i, row in enumerate(reader):
            if i == 0:
                rows.append("COLUMNS: " + " | ".join(row))
            else:
                rows.append(" | ".join(row))
    return "\n".join(rows)


def _clean_text(raw: str) -> str:
    """Light-touch cleaning: collapse whitespace runs, fix common OCR artifacts."""
    text = raw.replace("\x0c", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = text.replace("\x00", "")
    replacements = {"ﬁ": "fi", "ﬂ": "fl", "ﬀ": "ff", "ﬃ": "ffi", "ﬄ": "ffl"}
    for bad, good in replacements.items():
        text = text.replace(bad, good)
    return text.strip()


def _extract_structured_fields(text: str) -> dict[str, Any]:
    """Use Groq (LLaMA 3) to pull structured fields from raw legal text."""
    snippet = text[:6000]
    try:
        response = _client.chat.completions.create(
            model=_MODEL,
            max_tokens=1024,
            messages=[
                {
                    "role": "system",
                    "content": "You are a legal document analyst. Always respond with valid JSON only, no markdown fences.",
                },
                {
                    "role": "user",
                    "content": _STRUCT_PROMPT + snippet,
                },
            ],
        )
        raw_json = response.choices[0].message.content.strip()
        # Strip markdown fences if model adds them
        raw_json = re.sub(r"^```json?\s*", "", raw_json)
        raw_json = re.sub(r"\s*```$", "", raw_json)
        return json.loads(raw_json)
    except Exception as exc:
        return {"error": str(exc), "summary": "Could not parse structured fields."}


def process_document(file_path: str) -> dict[str, Any]:
    """
    Main entry point. Accepts a PDF, image, or text file.
    Returns a dict with:
      - doc_id: unique ID
      - file_name: original file name
      - raw_text: cleaned extracted text
      - structured_fields: LLM-extracted metadata
      - chunks: list of text chunks ready for embedding
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        raw_text = _extract_text_pdf(file_path)
    elif suffix in {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"}:
        raw_text = _extract_text_image(file_path)
    elif suffix == ".docx":
        raw_text = _extract_text_docx(file_path)
    elif suffix == ".rtf":
        raw_text = _extract_text_rtf(file_path)
    elif suffix == ".csv":
        raw_text = _extract_text_csv(file_path)
    else:
        # .txt, .md, .doc (plain fallback), .log, etc.
        raw_text = _extract_text_plain(file_path)

    clean_text = _clean_text(raw_text)
    structured_fields = _extract_structured_fields(clean_text)
    chunks = _chunk_text(clean_text)

    return {
        "doc_id": str(uuid.uuid4()),
        "file_name": path.name,
        "file_path": str(path.resolve()),
        "raw_text": clean_text,
        "structured_fields": structured_fields,
        "chunks": chunks,
        "char_count": len(clean_text),
        "chunk_count": len(chunks),
    }


def _chunk_text(text: str, chunk_size: int = 800, overlap: int = 150) -> list[dict]:
    """Split text into overlapping chunks for embedding."""
    words = text.split()
    chunks = []
    start = 0
    chunk_idx = 0

    while start < len(words):
        end = min(start + chunk_size, len(words))
        chunk_words = words[start:end]
        chunk_text = " ".join(chunk_words)
        chunks.append({
            "chunk_id": chunk_idx,
            "text": chunk_text,
            "word_start": start,
            "word_end": end,
        })
        chunk_idx += 1
        if end == len(words):
            break
        start += chunk_size - overlap

    return chunks
